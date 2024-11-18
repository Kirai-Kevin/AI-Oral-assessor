import streamlit as st
import openai
import speech_recognition as sr
import tempfile
from gtts import gTTS
import os
import time
from pathlib import Path
from playsound import playsound
from dotenv import load_dotenv
import json
from datetime import datetime
import docx
from docx.shared import Inches
import pandas as pd
import PyPDF2
import io
import sounddevice as sd

# Load environment variables
load_dotenv()

# Initialize OpenAI client with API key from .env
client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Initialize speech recognizer
recognizer = sr.Recognizer()

# Default rubric if none is provided
DEFAULT_RUBRIC = {
    "code_quality": {
        "weight": 0.3,
        "criteria": {
            "readability": 10,
            "organization": 10,
            "documentation": 10
        }
    },
    "functionality": {
        "weight": 0.4,
        "criteria": {
            "correctness": 10,
            "efficiency": 10,
            "error_handling": 10
        }
    },
    "best_practices": {
        "weight": 0.3,
        "criteria": {
            "naming_conventions": 10,
            "code_structure": 10,
            "standards_compliance": 10
        }
    }
}

def initialize_session_state():
    """Initialize session state variables"""
    if 'current_question' not in st.session_state:
        st.session_state.current_question = None
    if 'questions_asked' not in st.session_state:
        st.session_state.questions_asked = []
    if 'conversation_history' not in st.session_state:
        st.session_state.conversation_history = []
    if 'assessment_results' not in st.session_state:
        st.session_state.assessment_results = []
    if 'uploaded_files_content' not in st.session_state:
        st.session_state.uploaded_files_content = {}
    if 'rubric' not in st.session_state:
        st.session_state.rubric = DEFAULT_RUBRIC
    if 'submitted_code' not in st.session_state:
        st.session_state.submitted_code = None
    

def parse_rubric_file(uploaded_file):
    """Parse uploaded rubric file based on its format"""
    if uploaded_file is None:
        return DEFAULT_RUBRIC
    
    file_extension = Path(uploaded_file.name).suffix.lower()
    
    try:
        if file_extension == '.json':
            return json.loads(uploaded_file.getvalue().decode('utf-8'))
        elif file_extension == '.csv':
            df = pd.read_csv(uploaded_file)
            # Convert DataFrame to rubric format
            rubric = {}
            for _, row in df.iterrows():
                category = row['Category']
                if category not in rubric:
                    rubric[category] = {"weight": float(row['Weight']), "criteria": {}}
                rubric[category]["criteria"][row['Criterion']] = float(row['Max_Score'])
            return rubric
        elif file_extension == '.pdf':
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.getvalue()))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            # Parse PDF text into rubric format (implement based on expected format)
            # This is a simplified example
            return DEFAULT_RUBRIC
        elif file_extension == '.txt':
            # Parse text file into rubric format (implement based on expected format)
            return DEFAULT_RUBRIC
        else:
            st.warning(f"Unsupported rubric file format: {file_extension}. Using default rubric.")
            return DEFAULT_RUBRIC
    except Exception as e:
        st.error(f"Error parsing rubric file: {str(e)}. Using default rubric.")
        return DEFAULT_RUBRIC

def read_file_content(uploaded_file):
    """Read and return the content of an uploaded file"""
    if uploaded_file is None:
        return None
    
    try:
        # Get the file extension
        file_extension = Path(uploaded_file.name).suffix.lower()
        
        # Read content based on file type
        if file_extension in ['.py', '.txt', '.md', '.json', '.yaml', '.yml', '.css', '.html', '.js', 
                            '.java', '.cpp', '.c', '.cs', '.rb', '.php', '.swift', '.go', '.rs']:
            content = uploaded_file.getvalue().decode('utf-8')
        else:
            content = f"Unsupported file type: {file_extension}"
        
        return content
    except Exception as e:
        return f"Error reading file: {str(e)}"

# Keep existing functions: text_to_speech, play_audio, record_audio

def analyze_code(code, rubric, additional_files=None):
    """Analyze code and generate initial questions based on rubric"""
    context = f"Main code:\n{code}\n\n"
    if additional_files:
        context += "Additional files:\n"
        for filename, content in additional_files.items():
            context += f"\n{filename}:\n{content}\n"
    
    # Include rubric information in the prompt
    rubric_context = json.dumps(rubric, indent=2)
    
    prompt = f"""
    Analyze this code and related files using the following rubric:
    {rubric_context}
    
    Code to analyze:
    {context}
    
    Generate 2-3 relevant questions that:
    1. Test understanding of the code's functionality and structure
    2. Assess adherence to the rubric criteria
    3. Consider relationships between files (if applicable)
    4. Cover both basic and advanced concepts
    
    Format each question to align with specific rubric criteria.
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}]
        )
        questions = [q.strip() for q in response.choices[0].message.content.split('\n') if q.strip()]
        return questions
    except Exception as e:
        st.error(f"Error analyzing code: {str(e)}")
        return []
    
def analyze_rubric_and_code(code, rubric, additional_files=None):
    """
    Analyze code based on rubric criteria and generate targeted questions
    Returns both initial analysis and questions
    """
    context = f"Main code:\n{code}\n\n"
    if additional_files:
        context += "Additional files:\n"
        for filename, content in additional_files.items():
            context += f"\n{filename}:\n{content}\n"
    
    rubric_context = json.dumps(rubric, indent=2)
    
    # First, analyze the code against rubric criteria
    analysis_prompt = f"""
    Analyze this code using the provided rubric:
    {rubric_context}
    
    Code to analyze:
    {context}
    
    Provide a detailed analysis addressing each rubric category and criterion. Format your response as:
    CATEGORY|CRITERION|INITIAL_SCORE|COMMENTS
    
    For each criterion:
    1. Assess how well the code meets the criterion
    2. Identify specific areas for examination
    3. Note potential discussion points
    """
    
    try:
        analysis_response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": analysis_prompt}]
        )
        
        initial_analysis = analysis_response.choices[0].message.content.strip()
        
        # Generate questions based on analysis and rubric
        question_prompt = f"""
        Based on this code analysis:
        {initial_analysis}
        
        And using this rubric:
        {rubric_context}
        
        Generate questions that:
        1. Test understanding of code elements that scored lower in the analysis
        2. Verify knowledge of well-implemented features
        3. Explore specific rubric criteria
        4. Include follow-up questions for each main question
        
        For each category in the rubric, generate at least one question.
        Format each question as:
        CATEGORY|CRITERION|QUESTION|FOLLOW_UP|EXPECTED_POINTS
        """
        
        questions_response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": question_prompt}]
        )
        
        # Parse and structure the questions
        questions = []
        for line in questions_response.choices[0].message.content.strip().split('\n'):
            if '|' in line:
                category, criterion, question, follow_up, points = line.split('|')
                questions.append({
                    'category': category.strip(),
                    'criterion': criterion.strip(),
                    'question': question.strip(),
                    'follow_up': follow_up.strip(),
                    'points': float(points.strip())
                })
        
        return questions, initial_analysis
        
    except Exception as e:
        st.error(f"Error in code analysis: {str(e)}")
        return [], "Error in analysis"

def evaluate_answer_with_rubric(question_data, answer, code_context, initial_analysis, rubric, additional_files=None):
    """Enhanced evaluation function that considers rubric criteria and previous analysis"""
    context = f"Main code:\n{code_context}\n\n"
    if additional_files:
        context += "Additional files:\n"
        for filename, content in additional_files.items():
            context += f"\n{filename}:\n{content}\n"
    
    evaluation_prompt = f"""
    Question Category: {question_data['category']}
    Question Criterion: {question_data['criterion']}
    Maximum Points: {question_data['points']}
    
    Initial Analysis:
    {initial_analysis}
    
    Question: {question_data['question']}
    Student's Answer: {answer}
    
    Rubric:
    {json.dumps(rubric, indent=2)}
    
    Code Context:
    {context}
    
    Evaluate the answer considering:
    1. Alignment with rubric criterion
    2. Accuracy and completeness
    3. Initial code analysis
    4. Technical understanding
    
    If the answer is incomplete or incorrect, consider using this follow-up:
    {question_data['follow_up']}
    
    Respond EXACTLY in this format (each on a new line, no additional text):
    ASSESSMENT_TYPE
    EXPLANATION
    SCORE_NUMBER
    NEEDS_FOLLOWUP
    FOLLOWUP_QUESTION
    
    Where:
    - ASSESSMENT_TYPE must be exactly GOOD, NEEDS_IMPROVEMENT, or NEEDS_FOLLOWUP
    - EXPLANATION is your detailed evaluation
    - SCORE_NUMBER must be a number between 0 and {question_data['points']} (just the number, no text)
    - NEEDS_FOLLOWUP must be exactly YES or NO
    - FOLLOWUP_QUESTION is your follow-up question if needed, or NONE if not needed
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": evaluation_prompt}]
        )
        
        # Split response into lines and remove empty lines
        lines = [line.strip() for line in response.choices[0].message.content.split('\n') if line.strip()]
        
        if len(lines) < 5:
            raise ValueError("Incomplete response format")
        
        # Parse score with better error handling
        try:
            # First try simple float conversion
            score = float(lines[2])
        except ValueError:
            # If that fails, try to extract just the numeric part
            import re
            numeric_match = re.search(r'(\d+\.?\d*)', lines[2])
            if numeric_match:
                score = float(numeric_match.group(1))
            else:
                score = 0  # Default to 0 if no valid number found
                
        # Ensure score is within valid range
        max_points = float(question_data['points'])
        score = min(max(0, score), max_points)
            
        return {
            'assessment_type': lines[0],
            'explanation': lines[1],
            'score': score,
            'needs_followup': lines[3].upper() == 'YES',
            'followup_question': lines[4] if lines[4] != 'NONE' else None,
            'category': question_data['category'],
            'criterion': question_data['criterion']
        }
        
    except Exception as e:
        st.error(f"Error in evaluation: {str(e)}")
        return {
            'assessment_type': 'ERROR',
            'explanation': f'Error during evaluation: {str(e)}',
            'score': 0,
            'needs_followup': False,
            'followup_question': None,
            'category': question_data['category'],
            'criterion': question_data['criterion']
        }

def text_to_speech(text):
    """Convert text to speech and return the path to the audio file"""
    try:
        tts = gTTS(text=text, lang='en')
        
        # Create a temporary file with .mp3 extension
        with tempfile.NamedTemporaryFile(delete=False, suffix='.mp3') as fp:
            temp_filename = fp.name
            tts.save(temp_filename)
            return temp_filename
    except Exception as e:
        st.error(f"Error in text-to-speech conversion: {str(e)}")
        return None

def play_audio(file_path):
    """Play audio file using Streamlit's native audio player"""
    if file_path is None:
        return
        
    try:
        # Read the audio file
        with open(file_path, 'rb') as audio_file:
            audio_bytes = audio_file.read()
            
        # Create a container for the audio player
        audio_container = st.empty()
        
        # Play the audio using Streamlit's audio component
        audio_container.audio(audio_bytes, format='audio/mp3')
        
        # Remove temporary file after loading
        try:
            os.remove(file_path)
        except:
            pass
            
    except Exception as e:
        st.error(f"Error playing audio: {str(e)}")

def initialize_audio_system():
    """Initialize audio system with automatic device selection and retry logic"""
    try:
        import sounddevice as sd
        import speech_recognition as sr
        from time import sleep
        
        st.info("🎤 Initializing audio system...")
        
        # Maximum number of retries
        max_retries = 3
        retry_count = 0
        
        while retry_count < max_retries:
            try:
                # Get list of all audio devices
                devices = sd.query_devices()
                input_devices = []
                
                # Log available devices for debugging
                st.write("Available audio devices:")
                for i, device in enumerate(devices):
                    if device['max_input_channels'] > 0:
                        st.write(f"- {device['name']} (ID: {i})")
                        input_devices.append({
                            'index': i,
                            'name': device['name'],
                            'channels': device['max_input_channels'],
                            'default_samplerate': device['default_samplerate']
                        })
                
                if not input_devices:
                    raise Exception("No input devices found")
                
                # Try to find best device in this order:
                # 1. Default system microphone
                # 2. Built-in microphone
                # 3. First available microphone
                default_device = None
                for dev in input_devices:
                    name_lower = dev['name'].lower()
                    if "default" in name_lower:
                        default_device = dev
                        break
                    elif "built-in" in name_lower:
                        default_device = dev
                        break
                
                if not default_device and input_devices:
                    default_device = input_devices[0]
                
                if not default_device:
                    raise Exception("No suitable microphone found")
                
                # Configure audio settings
                sd.default.device = default_device['index']
                sd.default.samplerate = int(default_device['default_samplerate'])
                sd.default.channels = 1
                
                # Test microphone initialization
                recognizer = sr.Recognizer()
                mic = sr.Microphone(device_index=default_device['index'])
                
                with mic as source:
                    # Quick test recording
                    st.info("🎤 Testing microphone...")
                    recognizer.adjust_for_ambient_noise(source, duration=1)
                    
                    # Try a quick audio capture
                    audio = recognizer.listen(source, timeout=1.0, phrase_time_limit=1.0)
                    
                    st.success(f"✅ Microphone initialized successfully: {default_device['name']}")
                    return True, default_device['index']
                    
            except Exception as e:
                retry_count += 1
                if retry_count < max_retries:
                    st.warning(f"Attempt {retry_count}/{max_retries} failed. Retrying...")
                    sleep(1)  # Wait before retrying
                else:
                    st.error(f"Failed to initialize microphone after {max_retries} attempts: {str(e)}")
                    st.write("""
                    Troubleshooting steps:
                    1. Check if your microphone is properly connected
                    2. Ensure microphone permissions are granted to your browser
                    3. Try selecting a different audio input device in your system settings
                    4. Restart your browser and try again
                    """)
                    return False, None
                    
    except Exception as e:
        st.error(f"Error in audio system initialization: {str(e)}")
        return False, None


def record_audio_enhanced():
    """Enhanced audio recording with better error handling and user feedback"""
    if not hasattr(record_audio_enhanced, 'initialized'):
        success, device_index = initialize_audio_system()
        if not success:
            return None
        record_audio_enhanced.initialized = True
        record_audio_enhanced.device_index = device_index
    
    try:
        recognizer = sr.Recognizer()
        with sr.Microphone(device_index=record_audio_enhanced.device_index) as source:
            # Create a progress bar for ambient noise adjustment
            progress_bar = st.progress(0)
            for i in range(10):
                progress_bar.progress((i + 1) * 0.1)
                time.sleep(0.1)
            
            st.info("🎤 Adjusting for ambient noise...")
            recognizer.adjust_for_ambient_noise(source, duration=1)
            
            # Add a visual indicator that the system is listening
            with st.spinner("🎤 Listening... Speak now!"):
                try:
                    audio = recognizer.listen(source, timeout=10.0, phrase_time_limit=30.0)
                    st.info("Processing your answer...")
                    
                    try:
                        text = recognizer.recognize_google(audio)
                        st.success("✅ Successfully recorded and transcribed!")
                        return text
                    except sr.UnknownValueError:
                        st.error("Could not understand the audio. Please try speaking more clearly.")
                        return None
                    except sr.RequestError as e:
                        st.error(f"Error with speech recognition service: {str(e)}")
                        return None
                        
                except sr.WaitTimeoutError:
                    st.error("No speech detected within timeout period. Please try again.")
                    return None
                    
    except Exception as e:
        st.error(f"Error during audio recording: {str(e)}")
        record_audio_enhanced.initialized = False  # Reset initialization flag
        return None

def initialize_audio():
    """Initialize audio system with better error handling"""
    try:
        import sounddevice as sd
        import speech_recognition as sr
        
        # Check for working audio devices
        success, message, device_index = initialize_audio_system()
        
        if not success:
            st.error(message)
            return False
            
        # Configure default audio device
        sd.default.device = device_index
        sd.default.samplerate = 44100
        sd.default.channels = 1
        
        # Test microphone initialization
        recognizer = sr.Recognizer()
        mic = sr.Microphone(device_index=device_index)
        
        with mic as source:
            # Quick test recording
            try:
                recognizer.adjust_for_ambient_noise(source, duration=0.5)
                st.success(f"✅ {message}")
                return True
            except Exception as e:
                st.error(f"Failed to initialize microphone: {str(e)}")
                return False
                
    except Exception as e:
        st.error(f"Error setting up audio system: {str(e)}")
        return False


def evaluate_answer(question, answer, code_context, rubric, additional_files=None):
    """Evaluate student's answer using the provided rubric"""
    context = f"Main code:\n{code_context}\n\n"
    if additional_files:
        context += "Additional files:\n"
        for filename, content in additional_files.items():
            context += f"\n{filename}:\n{content}\n"
    
    rubric_context = json.dumps(rubric, indent=2)
    
    prompt = f"""
    Context: Student submitted this code and related files:
    {context}
    
    Rubric for evaluation:
    {rubric_context}
    
    Question asked: {question}
    Student's answer: {answer}
    
    Evaluate the answer considering:
    1. All rubric criteria and their weights
    2. Correctness and completeness
    3. Understanding of concepts
    4. Code relationships (if applicable)
    
    Provide your response in exactly this format:
    ASSESSMENT_TYPE|Detailed explanation based on rubric criteria|Follow-up question if needed|Score
    
    Where:
    - ASSESSMENT_TYPE is either GOOD or NEEDS_FOLLOWUP
    - Score is from 0-10 based on rubric weights
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}]
        )
        
        content = response.choices[0].message.content.strip()
        parts = content.split('|')
        
        if len(parts) != 4:
            return {
                'assessment': 'ERROR',
                'explanation': 'Invalid response format from evaluation',
                'followup': None,
                'score': 0
            }
            
        try:
            score = float(parts[3].strip())
        except (ValueError, TypeError):
            score = 0
            
        return {
            'assessment': parts[0].strip(),
            'explanation': parts[1].strip(),
            'followup': parts[2].strip() if parts[2].strip().lower() != 'none' else None,
            'score': score
        }
    except Exception as e:
        st.error(f"Error evaluating answer: {str(e)}")
        return {
            'assessment': 'ERROR',
            'explanation': 'There was an error evaluating the answer.',
            'followup': None,
            'score': 0
        }
    
def evaluate_code_directly(code, rubric, additional_files=None):
    """Evaluate code directly against rubric criteria without questions"""
    context = f"Main code:\n{code}\n\n"
    if additional_files:
        context += "Additional files:\n"
        for filename, content in additional_files.items():
            context += f"\n{filename}:\n{content}\n"
    
    prompt = f"""
    Analyze this code against the provided rubric:
    {json.dumps(rubric, indent=2)}
    
    Code to analyze:
    {context}
    
    For each rubric category and criterion, provide a direct score and explanation.
    Format your response exactly as:
    CATEGORY|CRITERION|SCORE|EXPLANATION
    
    Where:
    - SCORE is a number between 0-10
    - Each assessment should be on a new line
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}]
        )
        
        evaluations = []
        for line in response.choices[0].message.content.strip().split('\n'):
            if '|' in line:
                category, criterion, score, explanation = line.split('|')
                evaluations.append({
                    'category': category.strip(),
                    'criterion': criterion.strip(),
                    'score': float(score.strip()),
                    'explanation': explanation.strip()
                })
        return evaluations
    except Exception as e:
        st.error(f"Error in code evaluation: {str(e)}")
        return []

def generate_report():
    """Generate a comprehensive report including both code-based and question-based grades"""
    # Get direct code evaluation
    code_evaluations = evaluate_code_directly(
        st.session_state.submitted_code,
        st.session_state.rubric,
        st.session_state.uploaded_files_content
    )
    
    # Calculate code-based scores
    code_scores = {}
    code_explanations = {}
    for category in st.session_state.rubric:
        category_evals = [e for e in code_evaluations if e['category'] == category]
        if category_evals:
            code_scores[category] = sum(e['score'] for e in category_evals) / len(category_evals)
            code_explanations[category] = [e['explanation'] for e in category_evals]
    
    code_weighted_score = sum(
        code_scores.get(category, 0) * details['weight']
        for category, details in st.session_state.rubric.items()
    )
    
    # Calculate question-based scores
    question_scores = {}
    for result in st.session_state.assessment_results:
        category = result['question']['category']
        if category not in question_scores:
            question_scores[category] = []
        question_scores[category].append(result['evaluation']['score'])
    
    question_category_scores = {
        category: sum(scores) / len(scores)
        for category, scores in question_scores.items()
    }
    
    question_weighted_score = sum(
        question_category_scores.get(category, 0) * details['weight']
        for category, details in st.session_state.rubric.items()
    )
    
    def get_letter_grade(score):
        if score >= 9.0: return 'A', 'Excellent'
        elif score >= 8.0: return 'B', 'Good'
        elif score >= 7.0: return 'C', 'Satisfactory'
        elif score >= 6.0: return 'D', 'Needs Improvement'
        else: return 'F', 'Unsatisfactory'
    
    # Generate report
    doc = docx.Document()
    
    # Add styles
    code_style = doc.styles.add_style('CodeBlock', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    code_font = code_style.font
    code_font.name = 'Courier New'
    code_font.size = docx.shared.Pt(10)
    paragraph_format = code_style.paragraph_format
    paragraph_format.space_before = docx.shared.Pt(6)
    paragraph_format.space_after = docx.shared.Pt(6)
    paragraph_format.left_indent = docx.shared.Inches(0.5)
    
    # Title and date
    doc.add_heading("Code Assessment Report", 0)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Code-based evaluation section
    doc.add_heading("Direct Code Evaluation", 1)
    code_letter, code_description = get_letter_grade(code_weighted_score)
    grade_para = doc.add_paragraph()
    grade_para.add_run(f"Code-Based Grade: {code_letter}\n").bold = True
    grade_para.add_run(f"Numerical Score: {code_weighted_score:.2f}/10\n")
    grade_para.add_run(f"Assessment: {code_description}")
    
    # Add detailed code evaluation
    for category in st.session_state.rubric:
        if category in code_scores:
            doc.add_heading(f"{category} Evaluation", 2)
            doc.add_paragraph(f"Score: {code_scores[category]:.2f}/10")
            for explanation in code_explanations[category]:
                doc.add_paragraph(f"• {explanation}")
    
    # Question-based evaluation section
    doc.add_heading("Interview Question Evaluation", 1)
    question_letter, question_description = get_letter_grade(question_weighted_score)
    grade_para = doc.add_paragraph()
    grade_para.add_run(f"Question-Based Grade: {question_letter}\n").bold = True
    grade_para.add_run(f"Numerical Score: {question_weighted_score:.2f}/10\n")
    grade_para.add_run(f"Assessment: {question_description}")
    
    # Calculate final combined grade
    final_score = (code_weighted_score + question_weighted_score) / 2
    final_letter, final_description = get_letter_grade(final_score)
    
    # Add final grade section
    doc.add_heading("Final Combined Grade", 1)
    final_para = doc.add_paragraph()
    final_para.add_run(f"Final Grade: {final_letter}\n").bold = True
    final_para.add_run(f"Final Score: {final_score:.2f}/10\n")
    final_para.add_run(f"Overall Assessment: {final_description}")
    
    # Add existing sections (code, rubric, etc.)
    doc.add_heading("Submitted Code", 1)
    doc.add_paragraph(st.session_state.submitted_code, style='CodeBlock')
    
    if st.session_state.uploaded_files_content:
        doc.add_heading("Additional Files", 1)
        for filename, content in st.session_state.uploaded_files_content.items():
            doc.add_heading(filename, 2)
            doc.add_paragraph(content, style='CodeBlock')
    
    # Add detailed question responses
    doc.add_heading("Question Responses", 1)
    for i, result in enumerate(st.session_state.assessment_results, 1):
        doc.add_heading(f"Question {i}", 2)
        doc.add_paragraph(f"Question: {result['question']['question']}")
        doc.add_paragraph(f"Category: {result['question']['category']}")
        doc.add_paragraph(f"Student's Answer: {result['answer']}")
        doc.add_paragraph(f"Evaluation: {result['evaluation']['explanation']}")
        doc.add_paragraph(f"Score: {result['evaluation']['score']}/{result['question']['points']}")
    
    # Add recommendations
    doc.add_heading("Recommendations", 1)
    recommendations = generate_recommendations(final_score, code_scores)
    for rec in recommendations:
        doc.add_paragraph(f"• {rec}")
    
    # Save report
    report_filename = f"assessment_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(report_filename)
    return report_filename

def generate_recommendations(final_score, code_scores):
    """Generate recommendations based on both code evaluation and interview performance"""
    recommendations = []
    
    if final_score < 7.0:
        recommendations.append(
            "Focus on improving both code implementation and theoretical understanding of programming concepts."
        )
    
    for category, score in code_scores.items():
        if score < 7.0:
            if category == "code_quality":
                recommendations.append(
                    "Improve code readability, documentation, and organization. Consider using more descriptive "
                    "variable names and adding comprehensive comments."
                )
            elif category == "functionality":
                recommendations.append(
                    "Focus on implementing robust error handling, improving algorithmic efficiency, "
                    "and ensuring all requirements are met."
                )
            elif category == "best_practices":
                recommendations.append(
                    "Review and apply industry-standard coding conventions, design patterns, "
                    "and best practices for your programming language."
                )
    
    if not recommendations:
        recommendations.append(
            "Continue maintaining high coding standards while exploring advanced concepts "
            "and modern development practices."
        )
    
    return recommendations

def main():
    st.title("Directed's Code Assessment Interview")
    

    
    initialize_session_state()
    
    with st.sidebar:
        st.subheader("Instructions")
        st.write("""
        1. (Optional) Upload a custom grading rubric
        2. Submit your code (file upload or direct input)
        3. Upload any additional related files (optional)
        4. Click 'Start Assessment' to begin
        5. Listen to each question
        6. Click 'Record Answer' to respond verbally
        7. Review your assessment report
        """)
    
    # Rubric upload section
    st.subheader("Grading Rubric (Optional)")
    rubric_file = st.file_uploader(
        "Upload custom rubric (JSON, CSV, docx, doc, PDF, or TXT):",
        type=['json', 'csv', 'pdf', 'txt', 'docx', 'doc']
    )
    
    if rubric_file:
        st.session_state.rubric = parse_rubric_file(rubric_file)
        st.write("Custom rubric loaded successfully!")
    
    # Code submission section
    st.subheader("Code Submission")
    submission_method = st.radio(
        "Choose submission method:",
        ["Direct Input", "File Upload"]
    )
    
    if submission_method == "Direct Input":
        code = st.text_area("Enter your code here:", height=200)
        if code:
            st.session_state.submitted_code = code
    else:
        uploaded_code = st.file_uploader(
            "Upload your main code file:",
            type=['py', 'java', 'cpp', 'c', 'cs', 'rb', 'php', 'js', 'html', 'css', 'swift', 'go', 'rs', 'jsx']
        )
        if uploaded_code:
            st.session_state.submitted_code = read_file_content(uploaded_code)
    
    # Additional files section
    st.subheader("Additional Files (Optional)")
    uploaded_files = st.file_uploader(
        "Upload related files:",
        accept_multiple_files=True,
        type=['py', 'java', 'cpp', 'c', 'cs', 'rb', 'php', 'js', 'html', 'css', 'swift', 'go', 'rs', 'txt', 'md', 'json', 'yaml', 'yml']
    )
    
    if uploaded_files:
        for uploaded_file in uploaded_files:
            content = read_file_content(uploaded_file)
            if content:
                st.session_state.uploaded_files_content[uploaded_file.name] = content
    
    # Start assessment button
    if st.button("Start Assessment") and st.session_state.submitted_code:
        with st.spinner("Analyzing code and generating questions..."):
            # Get questions and initial analysis
            questions, analysis = analyze_rubric_and_code(
                st.session_state.submitted_code,
                st.session_state.rubric,
                st.session_state.uploaded_files_content
            )
            
            # Store in session state
            st.session_state.questions_asked = questions
            st.session_state.initial_analysis = analysis
            st.session_state.assessment_results = []
            
            if questions:
                st.session_state.current_question = 0
                st.rerun()
    
    # Assessment process
    if (st.session_state.submitted_code and 
        st.session_state.current_question is not None and 
        st.session_state.questions_asked):
        
        st.subheader("Assessment in Progress")
        
        # Show rubric category being assessed
        current_q = st.session_state.questions_asked[st.session_state.current_question]
        st.info(f"Category: {current_q['category']} - Criterion: {current_q['criterion']}")
        
        # Progress bar
        progress = st.progress((st.session_state.current_question + 1) / len(st.session_state.questions_asked))
        
        # Display current question
        st.write(f"Question {st.session_state.current_question + 1}: {current_q['question']}")
        
        # Text-to-speech
        with st.spinner("Speaking question..."):
            audio_file = text_to_speech(current_q['question'])
            play_audio(audio_file)
        
        # Record answer button
        if st.button("Record Answer"):
            answer = record_audio_enhanced()
            if answer:
                st.write(f"Your answer: {answer}")
                
                # Evaluate answer
                with st.spinner("Evaluating answer..."):
                    evaluation = evaluate_answer_with_rubric(
                        current_q,
                        answer,
                        st.session_state.submitted_code,
                        st.session_state.initial_analysis,
                        st.session_state.rubric,
                        st.session_state.uploaded_files_content
                    )
                
                if evaluation['assessment_type'] != 'ERROR':
                    st.write(f"Evaluation: {evaluation['explanation']}")
                    st.write(f"Score: {evaluation['score']}/{current_q['points']}")
                    
                    if evaluation['needs_followup']:
                        st.warning("Follow-up question needed:")
                        st.write(evaluation['followup_question'])
                        
                        # Add button for follow-up response
                        if st.button("Answer Follow-up"):
                            followup_answer = record_audio_enhanced()
                            if followup_answer:
                                st.write(f"Your follow-up answer: {followup_answer}")
                                # Re-evaluate with follow-up
                                evaluation = evaluate_answer_with_rubric(
                                    current_q,
                                    f"{answer}\nFollow-up answer: {followup_answer}",
                                    st.session_state.submitted_code,
                                    st.session_state.initial_analysis,
                                    st.session_state.rubric,
                                    st.session_state.uploaded_files_content
                                )
                                st.write(f"Updated evaluation: {evaluation['explanation']}")
                                st.write(f"Final score: {evaluation['score']}/{current_q['points']}")
                    
                    # Add result to assessment history
                    st.session_state.assessment_results.append({
                        'question': current_q,
                        'answer': answer,
                        'evaluation': evaluation
                    })
                    
                    # Move to next question or finish assessment
                    if st.session_state.current_question < len(st.session_state.questions_asked) - 1:
                        st.session_state.current_question += 1
                        st.rerun()
                    else:
                        st.success("Assessment completed!")
                        
                        # Generate and offer report download
                        with st.spinner("Generating assessment report..."):
                            report_file = generate_report()
                            
                            # Display summary
                            total_score = sum(result['evaluation']['score'] for result in st.session_state.assessment_results)
                            average_score = total_score / len(st.session_state.assessment_results)
                            
                            st.subheader("Assessment Summary")
                            st.write(f"Total Questions: {len(st.session_state.assessment_results)}")
                            st.write(f"Average Score: {average_score:.2f}/10")
                            
                            # Display scores by rubric category
                            st.subheader("Scores by Category")
                            for category, details in st.session_state.rubric.items():
                                category_score = average_score * details['weight']
                                st.write(f"{category}: {category_score:.2f}/10 (Weight: {details['weight']})")
                            
                            # Offer report download
                            with open(report_file, 'rb') as f:
                                st.download_button(
                                    label="Download Detailed Report",
                                    data=f,
                                    file_name=report_file,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                            
                            # Clean up
                            os.remove(report_file)
                        
                        # Reset session state for new assessment
                        st.session_state.current_question = None
                        
                        # Option to start new assessment
                        if st.button("Start New Assessment"):
                            st.session_state.clear()
                            st.rerun()

if __name__ == "__main__":
    main()